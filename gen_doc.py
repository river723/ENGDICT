# -*- coding: utf-8 -*-
"""生成《考研英语生词本》产品功能与变现方案 Word 文档。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

CN_FONT = "Microsoft YaHei"
BRAND = RGBColor(0x2E, 0x5C, 0x8A)
GRAY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# 默认字体（中英文）
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)


def set_cn(run, size=10.5, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def heading(text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.space_before = Pt(14)
        run = p.add_run(text)
        set_cn(run, 15, True, BRAND)
        pPr = p._p.get_or_add_pPr()
        pbdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pbdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "6",
            qn("w:space"): "2", qn("w:color"): "2E5C8A"})
        pbdr.append(bottom)
        pPr.append(pbdr)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    else:
        run = p.add_run(text)
        set_cn(run, 12, True, RGBColor(0x33, 0x44, 0x55))
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p


def body(text, size=10.5, color=None, after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn(run, size, False, color)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_cn(r1, 10.5, True)
        r2 = p.add_run(text)
        set_cn(r2, 10.5, False)
    else:
        r = p.add_run(text)
        set_cn(r, 10.5, False)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    return p


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        set_cn(run, 10, True, RGBColor(0xFF, 0xFF, 0xFF))
        shd = hdr[i]._tc.get_or_add_tcPr().makeelement(
            qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "2E5C8A"})
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            set_cn(run, 10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


# ============ 封面 ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("考研英语生词本")
set_cn(r, 30, True, BRAND)
title.paragraph_format.space_before = Pt(90)
title.paragraph_format.space_after = Pt(4)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("产品功能与商业变现方案")
set_cn(r, 16, False, GRAY)
sub.paragraph_format.space_after = Pt(50)

for line in [
    "纯本地运行 · 科学背词 · 4801 词增强词库",
    "跨平台：iOS / Android / Web / Windows 桌面",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    set_cn(r, 11.5, False, RGBColor(0x44, 0x44, 0x44))
    p.paragraph_format.space_after = Pt(4)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("\n\n版本 1.0.0 · 编制日期 2026-07-14")
set_cn(r, 10, False, GRAY)

doc.add_page_break()

# ============ 一、产品概述 ============
heading("一、产品概述", 1)
body("考研英语生词本是一款面向考研人群的英语词汇学习应用，主打完全离线、纯本地运行——"
     "无需联网、无需注册、无需任何 API Key，下载即用，隐私数据不出本机。"
     "应用内置一套经过真题精筛的 4801 词考研增强词库，涵盖释义、词根词缀、例句、"
     "易混词、记忆技巧与难度/考频标注，并配合艾宾浩斯遗忘曲线复习算法，"
     "帮助用户把「添加 → 学习 → 复习 → 掌握」形成科学的记忆闭环。")

body("应用采用 React Native + Expo 技术栈一套代码多端交付，可同时覆盖 "
     "iOS、Android、Web 以及 Windows 桌面（Electron 打包），"
     "为多渠道分发与变现提供了天然基础。")

heading("核心定位", 2)
bullet("目标人群：", "考研学生为主，兼顾四六级、专升本等应试英语人群。")
bullet("核心价值：", "离线、纯净、无广告干扰的高效背词工具 + 高质量精选词库。")
bullet("差异化：", "词库经近 20 年真题精筛，剔除超简单基础词与低频冷词，只保留有考察价值的词，配合词根词缀与记忆技巧，学得更少、记得更牢。")

# ============ 二、功能模块 ============
heading("二、功能模块详解", 1)
body("应用以底部四大 Tab（学习 / 生词本 / 词库 / 我的）为主干，覆盖从选词、"
     "学习、复习到统计与数据管理的完整学习链路。", after=8)

heading("1. 学习首页（智能仪表盘）", 2)
bullet("今日智能建议：", "根据当日计划、正确率与困难词情况，动态给出「继续学习 / 强化复习 / 添加生词」等个性化行动建议。")
bullet("今日进度：", "一行展示待学、已完成、正确率，并配进度条实时反馈。")
bullet("最近添加与一周趋势：", "横向瓦片快速回顾近期生词，折叠卡片展示本周学习词数与日均学习次数。")

heading("2. 多模式学习", 2)
bullet("单词卡（Flashcard）：", "点击翻转，正面单词与音标、背面释义/词根词缀/记忆技巧/例句，支持发音。")
bullet("听写模式：", "听单词发音拼写单词，训练音形对应。")
bullet("释义选择：", "看词选义的四选一测验，快速检验记忆。")
bullet("答错重试机制：", "答错的单词自动回到队列末尾，需连续答对才算过关，强化薄弱词。")

heading("3. 艾宾浩斯科学复习", 2)
bullet("遗忘曲线复习：", "按 1/2/4/7/15/30 天间隔自动安排复习计划，间隔可自定义。")
bullet("新词限量：", "每日新词数量可调（1–50），控制学习负荷。")
bullet("困难词强化：", "自动识别历史正确率偏低的单词，一键进入强化复习。")

heading("4. 内置词库浏览与查询", 2)
bullet("整库浏览：", "4801 词考研词库支持流畅滚动浏览。")
bullet("搜索：", "按单词或中文释义即时搜索，匹配优先级智能排序。")
bullet("排序与筛选：", "支持字母、难度、考频、乱序等排序，并可按难度、考频筛选。")
bullet("单词详情：", "释义、音标、词根词缀、易混词、记忆技巧、难度与考频完整呈现。")

heading("5. 生词本管理", 2)
bullet("从词库选词：", "勾选加入生词本，自动带出释义、词根、例句等增强信息。")
bullet("手工录入：", "自定义添加单词，自动用本地词库补全释义。")
bullet("难度与考频标注：", "每个单词标注 1–5 星难度与考频，辅助优先级判断。")

heading("6. 学习统计", 2)
bullet("今日概览：", "今日学习量与正确率。")
bullet("词汇统计：", "总词数、已掌握数（正确率 ≥ 80%）及掌握进度。")
bullet("一周趋势图：", "柱高代表学习量、颜色代表正确率、底部显示计划完成率。")
bullet("困难词与里程碑：", "列出需加强单词，并以里程碑激励持续学习。")

heading("7. 个性化设置与数据管理", 2)
bullet("外观主题：", "浅色 / 深色 / 跟随系统三种模式。")
bullet("显示开关：", "熟词僻义、词根词缀、记忆技巧可自由开关。")
bullet("发音设置：", "发音开关与自动发音，离线 TTS（移动端 expo-speech，Web 端浏览器语音）。")
bullet("数据备份：", "导出 / 导入 .bk 压缩备份文件，支持跨设备迁移与恢复。")

# ============ 三、产品特点 ============
heading("三、产品核心特点与优势", 1)
make_table(
    ["特点", "说明", "变现价值"],
    [
        ["完全离线", "本地词库 + 本地存储，断网可用，无服务器成本", "零运维成本，利润率高"],
        ["精选增强词库", "4801 词真题精筛，含词根词缀、记忆技巧、易混词", "核心付费卖点，可扩展多词库"],
        ["科学复习算法", "艾宾浩斯遗忘曲线 + 答错重试机制", "提升留存与学习效果口碑"],
        ["多端覆盖", "iOS / Android / Web / Windows 一套代码", "多渠道分发，扩大触达"],
        ["隐私友好", "数据不出本机，无需注册登录", "契合隐私趋势，降低使用门槛"],
        ["无广告纯净", "专注学习体验，无干扰", "适合做付费/订阅而非广告变现"],
    ],
    widths=[1.2, 3.2, 2.0],
)

# ============ 四、技术架构 ============
heading("四、技术架构", 1)
make_table(
    ["层面", "技术选型"],
    [
        ["框架", "React Native 0.83 + Expo SDK 55"],
        ["语言", "TypeScript"],
        ["UI 组件", "React Native Paper（Material Design）"],
        ["导航", "React Navigation（底部 Tab + Stack）"],
        ["本地存储", "AsyncStorage 持久化"],
        ["发音", "expo-speech（移动端）/ 浏览器语音合成（Web）"],
        ["桌面端", "Electron + electron-builder（Windows 安装包）"],
        ["数据", "内置 worddict.json 增强词库（4801 词）"],
    ],
    widths=[1.6, 4.8],
)

# ============ 五、变现策略 ============
heading("五、商业变现策略建议", 1)
body("基于「离线纯本地 + 高质量精选词库」的产品特性，建议采用以订阅/买断为主、"
     "多元组合的变现模型，避免破坏无广告的纯净体验。", after=8)

heading("1. 免费增值（Freemium）—— 推荐主模型", 2)
bullet("免费版：", "开放基础背词功能与部分词库（如高频核心 1000 词），建立口碑与用户基数。")
bullet("会员订阅：", "解锁完整 4801 词增强词库、全部学习模式、云备份、学习报告等高级功能。定价建议月付 12–18 元、年付 68–98 元。")
bullet("买断制：", "面向反感订阅的用户提供一次性解锁（建议 68–128 元），与订阅并行。")

heading("2. 词库内容付费（可持续扩展）", 2)
bullet("多词库商店：", "在现有词库架构（DICTIONARIES 注册表）基础上扩展四六级、考博、专四专八、雅思托福等词库，按包售卖。")
bullet("名师/机构联名词库：", "与考研名师或机构合作推出联名精选词库，分成变现。")

heading("3. 增值服务", 2)
bullet("云同步与多端备份：", "把当前 .bk 手动备份升级为账号云同步，作为付费点。")
bullet("学习报告与打卡社群：", "生成阶段性学习报告、排行榜、打卡激励，提升粘性与付费意愿。")
bullet("桌面版单独售卖：", "Windows 桌面版作为差异化产品面向长时间备考人群。")

heading("4. 渠道分发", 2)
bullet("移动端：", "App Store、各大安卓应用市场（应用宝、华为、小米、OPPO/vivo 等）上架。")
bullet("Web/桌面：", "官网直接分发下载，规避应用商店抽成，提升利润率。")

heading("变现路线图建议", 2)
make_table(
    ["阶段", "重点", "目标"],
    [
        ["第一阶段", "上架免费版 + 埋点，打磨体验与口碑", "积累种子用户与评价"],
        ["第二阶段", "上线会员订阅与买断，接入支付", "跑通付费闭环"],
        ["第三阶段", "扩展多词库商店、云同步与学习报告", "提升 ARPU 与留存"],
        ["第四阶段", "机构联名、多端联动、私域社群", "规模化增长"],
    ],
    widths=[1.2, 3.4, 1.8],
)

# ============ 六、风险与建议 ============
heading("六、落地注意事项", 1)
bullet("内容版权：", "词库释义、例句需确保自有或已获授权，避免版权风险，这是付费产品的底线。")
bullet("支付合规：", "iOS 内购须走 Apple 内购通道；安卓与 Web/桌面可用微信/支付宝，注意各平台分成规则。")
bullet("免费与付费的平衡：", "免费部分要足够好用以形成口碑，同时清晰体现付费版增量价值。")
bullet("数据资产化：", "在尊重隐私前提下，逐步引入可选账号体系，为云同步与个性化推荐打基础。")

# 页脚
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("考研英语生词本 · 产品功能与变现方案")
set_cn(fr, 8, False, GRAY)

out = "考研英语生词本-产品功能与变现方案.docx"
doc.save(out)
print("saved:", out)
